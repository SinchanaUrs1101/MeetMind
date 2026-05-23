from __future__ import annotations

import re
from datetime import datetime
from io import BytesIO
from typing import Any

from docx import Document
from docx.shared import Pt
from fpdf import FPDF


def _sanitize_filename(text: str) -> str:
    safe = re.sub(r"[^A-Za-z0-9 _-]", "", text or "meeting")
    safe = safe.strip().replace(" ", "_")
    return safe or "meeting"


def _normalize_text(value: Any) -> str:
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, dict):
        return str(value)
    return str(value or "")


def build_report_text(
    title: str,
    summary: str,
    decisions: list[str],
    action_items: list[dict[str, Any]],
    risks: list[dict[str, Any]],
    open_questions: list[str],
    filename: str,
    timestamp: str,
) -> str:
    lines = [
        f"Meeting Title: {title}",
        f"Report Filename: {filename}",
        f"Generated: {timestamp}",
        "",
        "Executive Summary:",
        summary or "No summary available.",
        "",
        "Decisions:",
    ]
    if decisions:
        lines.extend([f"- {line}" for line in decisions])
    else:
        lines.append("- None")

    lines.extend(["", "Action Items:"])
    if action_items:
        for item in action_items:
            lines.append(
                f"- {item.get('task', 'Untitled task')} | Owner: {item.get('owner') or 'Unassigned'} | Due: {item.get('due_date') or 'TBD'} | Status: {item.get('status') or 'pending'}"
            )
    else:
        lines.append("- None")

    lines.extend(["", "Risks / Dependencies:"])
    if risks:
        for risk in risks:
            lines.append(f"- {risk.get('risk_text', str(risk))}")
    else:
        lines.append("- None")

    lines.extend(["", "Open Questions:"])
    if open_questions:
        lines.extend([f"- {question}" for question in open_questions])
    else:
        lines.append("- None")

    return "\n".join(lines)


def create_meeting_pdf(
    title: str,
    summary: str,
    decisions: list[str],
    action_items: list[dict[str, Any]],
    risks: list[dict[str, Any]],
    open_questions: list[str],
    filename: str,
    timestamp: str,
) -> bytes:
    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", "B", 20)
    pdf.set_text_color(255, 255, 255)
    pdf.set_fill_color(67, 21, 34)
    page_width = pdf.w - pdf.l_margin - pdf.r_margin

    pdf.cell(page_width, 14, title, ln=True, fill=True)
    pdf.ln(5)

    def _safe_text(text: str) -> str:
        text = str(text or "")
        return re.sub(
            r"([^\s]{45})",
            lambda match: match.group(1) + "\u200b",
            text,
        )

    pdf.set_font("Arial", size=11)
    pdf.set_text_color(0, 0, 0)
    pdf.multi_cell(page_width, 6, _safe_text(f"Filename: {filename}"))
    pdf.multi_cell(page_width, 6, _safe_text(f"Generated: {timestamp}"))
    pdf.ln(6)

    def add_section(header: str, content: str | list[str]):
        pdf.set_font("Arial", "B", 14)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(page_width, 8, header, ln=True)
        pdf.ln(2)
        pdf.set_font("Arial", size=11)
        pdf.set_text_color(20, 20, 20)
        if isinstance(content, str):
            pdf.multi_cell(page_width, 7, _safe_text(content or "None"))
        else:
            if content:
                for line in content:
                    pdf.multi_cell(page_width, 7, _safe_text(f"- {line}"))
            else:
                pdf.multi_cell(page_width, 7, _safe_text("- None"))
        pdf.ln(6)

    add_section("Executive Summary", summary or "No summary available.")
    add_section("Decisions", decisions)
    add_section(
        "Action Items",
        [
            _safe_text(
                f"{item.get('task', 'Untitled task')} | Owner: {item.get('owner') or 'Unassigned'} | Due: {item.get('due_date') or 'TBD'} | Status: {item.get('status') or 'pending'}"
            )
            for item in action_items
        ],
    )
    add_section("Risks / Dependencies", [_safe_text(risk.get('risk_text', str(risk))) for risk in risks])
    add_section("Open Questions", [_safe_text(question) for question in open_questions])

    content = pdf.output(dest="S")
    if isinstance(content, str):
        return content.encode("latin-1")
    return bytes(content)


def create_meeting_docx(
    title: str,
    summary: str,
    decisions: list[str],
    action_items: list[dict[str, Any]],
    risks: list[dict[str, Any]],
    open_questions: list[str],
    filename: str,
    timestamp: str,
) -> bytes:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    document.add_heading(title, level=1)
    document.add_paragraph(f"Filename: {filename}")
    document.add_paragraph(f"Generated: {timestamp}")
    document.add_paragraph()

    document.add_heading("Executive Summary", level=2)
    document.add_paragraph(summary or "No summary available.")

    document.add_heading("Decisions", level=2)
    if decisions:
        for decision in decisions:
            document.add_paragraph(_normalize_text(decision), style="List Bullet")
    else:
        document.add_paragraph("None")

    document.add_heading("Action Items", level=2)
    if action_items:
        for item in action_items:
            document.add_paragraph(
                _normalize_text(item.get("task"))
                + f" | Owner: {_normalize_text(item.get('owner') or 'Unassigned')} | Due: {_normalize_text(item.get('due_date') or 'TBD')} | Status: {_normalize_text(item.get('status') or 'pending')}",
                style="List Bullet",
            )
    else:
        document.add_paragraph("None")

    document.add_heading("Risks / Dependencies", level=2)
    if risks:
        for risk in risks:
            document.add_paragraph(_normalize_text(risk.get("risk_text", str(risk))), style="List Bullet")
    else:
        document.add_paragraph("None")

    document.add_heading("Open Questions", level=2)
    if open_questions:
        for question in open_questions:
            document.add_paragraph(_normalize_text(question), style="List Bullet")
    else:
        document.add_paragraph("None")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_export_filename(title: str, extension: str) -> str:
    clean_title = _sanitize_filename(title)
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    return f"{clean_title}_{timestamp}.{extension}", timestamp
