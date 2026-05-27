import io
import json
from datetime import datetime

import streamlit as st
import streamlit.components.v1 as components
from utils.api_client import api_client

# =====================================================
# PAGE CONFIG
# =====================================================
st.set_page_config(
    page_title="MeetMind AI",
    page_icon="🧠",
    layout="wide",
)

# =====================================================
# HELPER — marquee logos HTML
# =====================================================
def _build_logos() -> str:
    logos = [
        ("S", "#6366f1", "Streamlit"),
        ("F", "#a855f7", "FastAPI"),
        ("P", "#8b5cf6", "PostgreSQL"),
        ("G", "#6366f1", "Groq"),
        
    ]
    html = ""
    for letter, color, name in logos:
        html += f"""
        <div class="marquee-logo">
          <div class="logo-icon liquid-glass" style="background:{color}33;">{letter}</div>
          <span class="logo-name">{name}</span>
        </div>"""
    return html


# =====================================================
# HELPER — export to .docx
# =====================================================
def meetings_to_docx(meetings: list) -> bytes:
    try:
        from docx import Document as DocxDocument
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = DocxDocument()
        t = doc.add_heading("MeetMind AI — Meeting Export", level=0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ts = doc.add_paragraph(f"Exported: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
        ts.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")

        for i, meeting in enumerate(meetings, start=1):
            doc.add_heading(f"Meeting {i}: {meeting.get('title', 'Untitled')}", level=1)
            for field in ["date", "attendees", "duration"]:
                val = meeting.get(field)
                if val:
                    p = doc.add_paragraph()
                    p.add_run(f"{field.capitalize()}: ").bold = True
                    p.add_run(str(val) if not isinstance(val, list) else ", ".join(val))
            if meeting.get("summary"):
                doc.add_heading("Summary", level=2)
                doc.add_paragraph(meeting["summary"])
            for items, key, label in [
                (meeting.get("action_items",   []),        "task",     "Action Items"),
                (meeting.get("decisions",      []),        "decision", "Decisions"),
                (meeting.get("risks",          []),        "risk",     "Risks"),
                (meeting.get("open_questions", []) or [], "question", "Open Questions"),
            ]:
                if items:
                    doc.add_heading(label, level=2)
                    for item in items:
                        doc.add_paragraph(
                            item.get(key) or item.get("description") or str(item),
                            style="List Bullet"
                        )
            if i < len(meetings):
                doc.add_paragraph("─" * 60)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    except ImportError:
        return json.dumps(meetings, default=str, indent=2).encode()


# =====================================================
# API DATA
# =====================================================
health_status     = api_client.health_check()
meetings_response = api_client.get_meetings()
meetings          = meetings_response if isinstance(meetings_response, list) else []

action_items   = [i for m in meetings for i in m.get("action_items",   [])]
decisions      = [i for m in meetings for i in m.get("decisions",      [])]
risks          = [i for m in meetings for i in m.get("risks",          [])]
open_questions = [q for m in meetings for q in (m.get("open_questions") or [])]
owners         = sorted({i.get("owner") for i in action_items if i.get("owner")})

total_meetings  = len(meetings)
total_actions   = len(action_items)
total_decisions = len(decisions)
total_risks_q   = len(risks) + len(open_questions)

# =====================================================
# GLOBAL CSS  (sidebar + content below hero)
# =====================================================
st.markdown("""
<style>
@import url('https://api.fontshare.com/v2/css?f[]=general-sans@400,500,600,700&display=swap');
@import url('https://fonts.googleapis.com/css2?family=Geist+Sans:wght@300;400;500;600;700&display=swap');

:root {
  --bg:       hsl(260, 87%, 3%);
  --fg:       hsl(40, 6%, 95%);
  --hero-sub: hsl(40, 6%, 82%);
}

html, body, [class*="css"] {
  font-family: 'Geist Sans', sans-serif !important;
  background: var(--bg) !important;
  color: var(--fg) !important;
}

#MainMenu { visibility: hidden; }
footer    { visibility: hidden; }
header    { visibility: hidden; }

.stApp { background: var(--bg) !important; }

.block-container {
  padding: 0 !important;
  max-width: 100% !important;
}

section[data-testid="stSidebar"] {
  background: hsl(260, 70%, 5%) !important;
  border-right: 1px solid rgba(255,255,255,0.06);
}
section[data-testid="stSidebar"] * { color: var(--fg) !important; }

.content-wrap {
  padding: 40px 48px;
  max-width: 1400px;
  margin: 0 auto;
}

.glass-card {
  background: rgba(255,255,255,0.04);
  border: 1px solid rgba(255,255,255,0.08);
  backdrop-filter: blur(14px);
  border-radius: 20px;
  padding: 28px;
  box-shadow: 0 8px 32px rgba(0,0,0,0.5);
  transition: transform 0.3s ease;
  color: var(--fg);
}
.glass-card:hover { transform: translateY(-4px); }
.glass-card h2, .glass-card h3 { color: var(--fg); margin-top: 0; }
.glass-card p  { color: var(--hero-sub); margin: 0; }

.progress-bar {
  width: 100%; height: 12px;
  border-radius: 999px;
  background: rgba(255,255,255,0.07);
  overflow: hidden;
}
.progress-fill {
  width: 200%; height: 100%;
  background: linear-gradient(90deg, #6366f1, #a855f7, #fcd34d, #6366f1);
  animation: slide 2.8s linear infinite;
}
@keyframes slide {
  0%   { transform: translateX(-50%); }
  100% { transform: translateX(0%); }
}

.stButton > button,
.stDownloadButton > button {
  width: 100% !important;
  border: 1px solid rgba(255,255,255,0.14) !important;
  border-radius: 999px !important;
  background: rgba(255,255,255,0.06) !important;
  color: var(--fg) !important;
  font-weight: 500 !important;
  padding: 0.65rem 1.2rem !important;
  font-family: 'Geist Sans', sans-serif !important;
  transition: background 0.2s ease, transform 0.2s ease !important;
}
.stButton > button:hover,
.stDownloadButton > button:hover {
  background: rgba(255,255,255,0.12) !important;
  transform: scale(1.02) !important;
}
</style>
""", unsafe_allow_html=True)

# =====================================================
# SIDEBAR
# =====================================================
with st.sidebar:
    st.markdown("## 🧠 MeetMind AI")
    st.markdown("---")
    st.markdown("**Navigation**")
    st.markdown("""
- 📹 Upload Meeting
- 📋 View Minutes
- ✅ Action Items
- 📊 Dashboard
""")
    st.markdown("---")
    st.markdown("**API Status**")
    if health_status.get("status") == "healthy":
        st.success("Backend Connected")
    else:
        st.error("Backend Offline")

# =====================================================
# HERO  (full-screen iframe — video, navbar, CTA, marquee)
# =====================================================
LOGOS_HTML = _build_logos()

components.html(f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8"/>
<link href="https://api.fontshare.com/v2/css?f[]=general-sans@400,500,600,700&display=swap" rel="stylesheet"/>
<link href="https://fonts.googleapis.com/css2?family=Geist+Sans:wght@300;400;500;600;700&display=swap" rel="stylesheet"/>
<style>
*,*::before,*::after{{box-sizing:border-box;margin:0;padding:0}}
:root{{--bg:hsl(260,87%,3%);--fg:hsl(40,6%,95%);--sub:hsl(40,6%,82%)}}
html,body{{height:100%;background:var(--bg);color:var(--fg);font-family:'Geist Sans',sans-serif;overflow:hidden}}

.hero-root{{position:relative;width:100vw;height:100vh;min-height:620px;display:flex;flex-direction:column;overflow:hidden}}

.bg-video{{position:absolute;inset:0;width:100%;height:100%;object-fit:cover;opacity:0;z-index:0}}

.blur-blob{{position:absolute;top:50%;left:50%;transform:translate(-50%,-50%);width:984px;height:527px;background:hsl(260,40%,4%);filter:blur(82px);opacity:.9;pointer-events:none;z-index:1}}

.content-layer{{position:relative;z-index:2;display:flex;flex-direction:column;height:100%}}

/* Navbar */
.navbar{{display:flex;align-items:center;justify-content:space-between;padding:20px 32px}}
.nav-logo{{font-family:'General Sans',sans-serif;font-size:19px;font-weight:700;color:var(--fg);display:flex;align-items:center;gap:9px}}
.logo-badge{{width:32px;height:32px;background:linear-gradient(135deg,#6366f1,#a855f7);border-radius:8px;display:flex;align-items:center;justify-content:center;font-size:16px}}
.nav-links{{display:flex;gap:2px;align-items:center}}
.nav-link{{background:none;border:none;cursor:pointer;color:rgba(255,255,255,.82);font-size:14px;font-weight:500;padding:8px 14px;border-radius:999px;font-family:'Geist Sans',sans-serif;display:flex;align-items:center;gap:5px;transition:background .18s}}
.nav-link:hover{{background:rgba(255,255,255,.08)}}
.chev{{font-size:9px;opacity:.55}}
.nav-cta{{background:rgba(255,255,255,.07);border:1px solid rgba(255,255,255,.15);border-radius:999px;color:var(--fg);font-size:14px;font-weight:500;padding:8px 20px;cursor:pointer;font-family:'Geist Sans',sans-serif;transition:background .18s}}
.nav-cta:hover{{background:rgba(255,255,255,.14)}}
.divider{{height:1px;background:linear-gradient(90deg,transparent,rgba(255,255,255,.16),transparent);margin-top:3px}}

/* Body */
.hero-body{{flex:1;display:flex;align-items:center;justify-content:center;flex-direction:column;text-align:center;padding:0 24px;overflow:visible}}
.headline{{font-family:'General Sans',sans-serif;font-size:clamp(72px,13vw,200px);font-weight:500;line-height:1.02;letter-spacing:-.024em;color:var(--fg);white-space:nowrap}}
.grad{{background:linear-gradient(to left,#6366f1,#a855f7,#fcd34d);-webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text}}
.subtitle{{color:var(--sub);font-size:17px;line-height:1.75;max-width:420px;margin-top:9px;opacity:.82}}
.hero-cta{{margin-top:25px;background:rgba(255,255,255,.07);border:1px solid rgba(255,255,255,.16);border-radius:999px;color:var(--fg);font-size:15px;font-weight:500;padding:14px 29px;cursor:pointer;font-family:'Geist Sans',sans-serif;letter-spacing:.01em;transition:background .2s,transform .2s}}
.hero-cta:hover{{background:rgba(255,255,255,.13);transform:scale(1.03)}}

/* Liquid glass */
.liquid-glass{{background:rgba(255,255,255,.01);backdrop-filter:blur(4px);-webkit-backdrop-filter:blur(4px);box-shadow:inset 0 1px 1px rgba(255,255,255,.1);position:relative;overflow:hidden}}
.liquid-glass::before{{content:"";position:absolute;inset:0;border-radius:inherit;padding:1.4px;background:linear-gradient(180deg,rgba(255,255,255,.45) 0%,rgba(255,255,255,.15) 20%,rgba(255,255,255,0) 40%,rgba(255,255,255,0) 60%,rgba(255,255,255,.15) 80%,rgba(255,255,255,.45) 100%);-webkit-mask:linear-gradient(#fff 0 0) content-box,linear-gradient(#fff 0 0);-webkit-mask-composite:xor;mask-composite:exclude;pointer-events:none}}

/* Marquee */
.marquee-section{{padding-bottom:34px;display:flex;align-items:center;justify-content:center;gap:48px;overflow:hidden;width:100%;max-width:860px;margin:0 auto}}
.marquee-label{{font-size:13px;color:rgba(255,255,255,.42);line-height:1.6;white-space:nowrap;flex-shrink:0}}
.marquee-wrap{{flex:1;overflow:hidden;mask-image:linear-gradient(90deg,transparent,black 10%,black 90%,transparent);-webkit-mask-image:linear-gradient(90deg,transparent,black 10%,black 90%,transparent)}}
.marquee-track{{display:flex;gap:44px;width:max-content;animation:mq 20s linear infinite}}
@keyframes mq{{0%{{transform:translateX(0%)}}100%{{transform:translateX(-50%)}}}}
.marquee-logo{{display:flex;align-items:center;gap:10px;flex-shrink:0}}
.logo-icon{{width:26px;height:26px;border-radius:7px;display:flex;align-items:center;justify-content:center;font-size:11px;font-weight:700;color:var(--fg)}}
.logo-name{{font-size:15px;font-weight:600;color:var(--fg);white-space:nowrap}}
</style>
</head>
<body>
<div class="hero-root">
  <video class="bg-video" id="bgvid"
    src="https://d8j0ntlcm91z4.cloudfront.net/user_38xzZboKViGWJOttwIXH07lWA1P/hf_20260328_065045_c44942da-53c6-4804-b734-f9e07fc22e08.mp4"
    muted playsinline preload="auto"></video>

  <div class="blur-blob"></div>

  <div class="content-layer">
    <div class="navbar">
      <div class="nav-logo">
        <div class="logo-badge">Ⓜ️</div>
        MeetMind AI
      </div>
      
    </div>
    <div class="divider"></div>

    <div class="hero-body">
      <div class="headline">Meet<span class="grad">Mind AI</span></div>
      <p class="subtitle">The most powerful AI ever deployed<br/>in meeting intelligence</p>
    </div>

    <div class="marquee-section">
      <div class="marquee-label">Technical<br/>Stack</div>
      <div class="marquee-wrap">
        <div class="marquee-track">
          {LOGOS_HTML}
          {LOGOS_HTML}
        </div>
      </div>
    </div>
  </div>
</div>

<script>
(function(){{
  var v=document.getElementById('bgvid'),FADE=500,LEAD=0.55,raf,out=false;
  function fadeTo(t,dur){{
    cancelAnimationFrame(raf);
    var s=performance.now(),f=parseFloat(v.style.opacity)||0;
    (function step(n){{
      var p=Math.min((n-s)/dur,1);
      v.style.opacity=f+(t-f)*p;
      if(p<1)raf=requestAnimationFrame(step);
    }})(performance.now());
  }}
  v.addEventListener('loadeddata',function(){{v.style.opacity=0;v.play().catch(function(){{}});fadeTo(1,FADE);}});
  v.addEventListener('timeupdate',function(){{
    var r=v.duration-v.currentTime;
    if(!out&&r<=LEAD&&r>0){{out=true;fadeTo(0,FADE);}}
  }});
  v.addEventListener('ended',function(){{
    v.style.opacity=0;
    setTimeout(function(){{v.currentTime=0;v.play().catch(function(){{}});out=false;fadeTo(1,FADE);}},100);
  }});
}})();
</script>
</body>
</html>""", height=720)

# =====================================================
# DASHBOARD CONTENT  (below hero)
# =====================================================
st.markdown('<div class="content-wrap">', unsafe_allow_html=True)

# ── Metric Cards ─────────────────────────────────────────────────────────
components.html(f"""<!DOCTYPE html><html><head>
<link href="https://fonts.googleapis.com/css2?family=Geist+Sans:wght@400;500;700&display=swap" rel="stylesheet"/>
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{background:transparent;font-family:'Geist Sans',sans-serif}}
.row{{display:flex;gap:18px}}
.card{{flex:1;border-radius:20px;padding:26px 22px;color:hsl(40,6%,95%);box-shadow:0 12px 32px rgba(0,0,0,.6);transition:transform .28s ease;min-width:0}}
.card:hover{{transform:translateY(-6px)}}
.label{{font-size:11px;font-weight:500;letter-spacing:.7px;text-transform:uppercase;opacity:.78}}
.value{{font-size:50px;font-weight:700;line-height:1;margin:10px 0 6px}}
.desc{{font-size:12px;opacity:.62}}
</style></head><body>
<div class="row">
  <div class="card" style="background:linear-gradient(135deg,#3730a3,#6366f1)">
    <div class="label">Total Meetings</div><div class="value">{total_meetings}</div><div class="desc">Meetings analyzed</div>
  </div>
  <div class="card" style="background:linear-gradient(135deg,#0369a1,#3b82f6)">
    <div class="label">Action Items</div><div class="value">{total_actions}</div><div class="desc">Tasks extracted</div>
  </div>
  <div class="card" style="background:linear-gradient(135deg,#b45309,#f59e0b)">
    <div class="label">Decisions</div><div class="value">{total_decisions}</div><div class="desc">Key decisions captured</div>
  </div>
  <div class="card" style="background:linear-gradient(135deg,#7e22ce,#a855f7)">
    <div class="label">Risks &amp; Questions</div><div class="value">{total_risks_q}</div><div class="desc">Pending blockers</div>
  </div>
</div>
</body></html>""", height=148)

st.markdown("<br>", unsafe_allow_html=True)

# ── Analytics + Quick Insights ───────────────────────────────────────────
left, right = st.columns([3, 1.4])

with left:
    st.markdown("""
    <div class="glass-card">
        <h2 style="margin-bottom:10px;">📈 Meeting Analytics</h2>
        <p style="margin-bottom:20px;">
            Visualize collaboration efficiency and meeting insights
            in real-time using AI-powered analysis.
        </p>
        <div class="progress-bar"><div class="progress-fill"></div></div>
        <p style="margin-top:16px;font-size:13px;opacity:0.55;">
            Keep uploading meeting transcripts to unlock richer analytics and insights.
        </p>
    </div>
    """, unsafe_allow_html=True)

with right:
    st.markdown(f"""
    <div class="glass-card">
        <h3 style="margin-bottom:16px;">⚡ Quick Insights</h3>
        <p style="margin-bottom:10px;">
            <b style="color:hsl(40,6%,95%);">👥 Owners Engaged:</b> {len(owners)}
        </p>
        <p style="margin-bottom:10px;">
            <b style="color:hsl(40,6%,95%);">🟢 AI Health:</b>
            {health_status.get('status', 'unknown')}
        </p>
        <p>
            <b style="color:hsl(40,6%,95%);">🕒 Last Sync:</b><br/>
            {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}
        </p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)


st.markdown("<br>", unsafe_allow_html=True)

# ── Getting Started ───────────────────────────────────────────────────────
st.markdown("""
<div class="glass-card">
    <h2 style="margin-bottom:16px;">🚀 Getting Started</h2>
    <p style="margin-bottom:10px;">1️⃣ Upload meeting transcripts or notes.</p>
    <p style="margin-bottom:10px;">
        2️⃣ AI automatically extracts summaries, action items, risks, and decisions.
    </p>
    <p>3️⃣ Review analytics and collaborate with your team efficiently.</p>
</div>
""", unsafe_allow_html=True)

st.markdown('</div>', unsafe_allow_html=True)
